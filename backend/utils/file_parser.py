import io
from typing import Optional


def parse_file(content: bytes, extension: str) -> str:
    ext = extension.lower().lstrip(".")
    if ext == "txt":
        return _parse_txt(content)
    if ext == "pdf":
        return _parse_pdf(content)
    if ext == "docx":
        return _parse_docx(content)
    raise ValueError(f"Unsupported file type: {ext}")


def _parse_txt(content: bytes) -> str:
    for encoding in ("utf-8", "utf-16", "latin-1", "cp1252"):
        try:
            return content.decode(encoding)
        except (UnicodeDecodeError, LookupError):
            continue
    return content.decode("utf-8", errors="replace")


def _parse_pdf(content: bytes) -> str:
    try:
        import pdfplumber
        text_parts = []
        with pdfplumber.open(io.BytesIO(content)) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
        return "\n".join(text_parts)
    except ImportError:
        return "[PDF parsing requires pdfplumber. Run: pip install pdfplumber]"
    except Exception as e:
        return f"[PDF parsing failed: {e}]"


def _parse_docx(content: bytes) -> str:
    try:
        from docx import Document
        doc = Document(io.BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except ImportError:
        return "[DOCX parsing requires python-docx. Run: pip install python-docx]"
    except Exception as e:
        return f"[DOCX parsing failed: {e}]"
